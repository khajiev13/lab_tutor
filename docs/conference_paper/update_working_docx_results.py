from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/khajievroma/Projects/lab_tutor/docs/conference_paper")
DOCX_PATH = ROOT / "lab_tutor_conference_paper_working.docx"
ARTIFACT_PATH = (
    ROOT.parents[1]
    / "backend"
    / "app"
    / "modules"
    / "marketdemandanalyst"
    / "notebooks"
    / "artifacts"
    / "big_data_employability_course_2_artifacts.json"
)
CONCLUSION_RESULTS_PREFIX = "The Big Data case study shows that curriculum enrichment improves estimated curriculum-to-market alignment"
ABSTRACT_PREFIXES = ("Abstract:",)
DATASET_SUMMARY_PREFIXES = (
    "The job dataset was collected from graph-linked search results",
)
EVALUATION_METHOD_PREFIXES = (
    "The evaluation compares the four curriculum variants against held-out market requirements.",
    "The evaluation now has two complementary parts.",
)
CASE_STUDY_QUESTION_PREFIXES = (
    "The case study is designed to answer two core questions.",
    "The case study is designed to answer two linked questions.",
)
CURRENT_ABSTRACT_PARAGRAPH = (
    "Abstract: Higher education courses increasingly face three linked challenges: teaching knowledge is fragmented "
    "across heterogeneous course materials, curricula evolve more slowly than labor-market skill demands, and the "
    "manual construction of supporting resources imposes substantial burden on instructors. To address these issues, "
    "we propose Lab Tutor, a knowledge-graph-centric multi-agent framework for intelligent curriculum resource "
    "construction and curriculum-market alignment. The framework organizes teacher-provided materials into a course "
    "knowledge graph, enriches that graph through textbook discovery and chapter-level skill extraction, aligns "
    "curriculum content with market-demand skills derived from job postings, and links identified skills to downstream "
    "learning resources under human-in-the-loop review checkpoints. We evaluate the framework through a Big Data "
    "course case study using two complementary analyses built on the same four curriculum variants: KG, KG + B_S, "
    "KG + J_S, and KG + B_S + J_S. In the course-level held-out benchmark, the strongest variant (KG + J_S) improved "
    "demand-weighted skill coverage from 0.2961 to 0.4806 and average job-fit from 0.2775 to 0.4602 across 8 held-out "
    "jobs drawn from 28 strict-scope Big Data roles. In a student-centered spillover analysis, skills extracted from "
    "7 student-selected target job postings improved average job-fit on the other 55 jobs from 0.2588 to 0.4259, "
    "while the fully personalized variant reached 0.4333. These findings suggest that a knowledge-graph-centric "
    "multi-agent workflow can improve estimated curriculum-to-market alignment and can also help explain how skills "
    "learned for a small target job portfolio transfer to adjacent opportunities, while preserving instructor oversight "
    "at critical decision points."
)
CURRENT_DATASET_SUMMARY_PARAGRAPH = (
    "The job dataset was collected from graph-linked search results and deduplicated into a broad pool of 62 jobs. "
    "For the main course-level benchmark, 28 strict-scope Big Data / Data Engineering jobs were retained and split "
    "into 20 build jobs and 8 held-out jobs using a seeded 0.30 held-out ratio. Because persisted descriptions are "
    "now available for the evaluated jobs, the notebook extracts job skills from raw job text rather than relying on "
    "graph-linked proxy skills. The paper therefore reports a small held-out benchmark based on raw-description "
    "extraction."
)
CURRENT_EVALUATION_METHOD_PARAGRAPH = (
    "The evaluation now has two complementary parts. The course-level benchmark compares the four curriculum variants "
    "against held-out Big Data job requirements using skills extracted from raw descriptions. The student-centered "
    "case study uses 7 job postings explicitly selected by STUDENT(id=2) as a seed portfolio and evaluates how that "
    "extracted skill bundle transfers to the remaining 55 jobs in the broad graph. Each skill instance is then judged "
    "as covered, partial, or missing relative to a curriculum variant, and these labels are mapped to numeric values "
    "of 1.0, 0.5, and 0.0, respectively."
)
CURRENT_CASE_STUDY_QUESTION_PARAGRAPH = (
    "The case study is designed to answer two linked questions. First, how much do course-document, book-skill, and "
    "market-skill enrichment improve held-out curriculum-to-market alignment for Big Data roles? Second, if a student "
    "prepares around 7 explicit target jobs, how far do those learned skills transfer to related jobs outside the "
    "seed portfolio? Framing the study in this way keeps the course-level benchmark separate from the personalized "
    "spillover analysis and makes the latest artifact easier to interpret."
)
COURSE_CHAPTER_COUNT = 6
UPLOADED_TRANSCRIPT_DOCUMENTS = 38
APPROVED_BOOKS = 3
CONCLUSION_RESULTS_PARAGRAPH = (
    "The Big Data case study shows that curriculum enrichment improves estimated curriculum-to-market alignment "
    "relative to a course-only baseline. In the current course-level held-out evaluation, market-derived skills "
    "contribute most of the observed lift, while textbook-derived skills provide additional but smaller gains. "
    "The personalized case study adds a second perspective: the student's 7 target postings span architecture, "
    "data engineering, full-stack, and junior software roles, and the extracted seed-job skills substantially "
    "improve fit for many other jobs in the graph. In practice, that transfer is carried by overlapping "
    "competencies such as SQL and data profiling, pipeline engineering, Python-based implementation, cloud-native "
    "tooling, APIs and microservices, and testing or automation. These findings do not demonstrate employment "
    "outcomes, but they do show that a knowledge-graph-centric multi-agent workflow can expose and reduce "
    "curriculum-market skill gaps in a measurable way while also illustrating how a student's chosen target jobs "
    "can support broader opportunity spillover."
)


def load_artifact() -> dict:
    if not ARTIFACT_PATH.exists():
        raise FileNotFoundError(ARTIFACT_PATH)
    return json.loads(ARTIFACT_PATH.read_text())


def success_count(summary: dict, key: str) -> str:
    count = int(round(summary[key] * summary["job_count"]))
    return f"{count}/{summary['job_count']}"


def format_row(variant: str, summary: dict) -> tuple[str, str, str, str, str]:
    return (
        variant,
        f"{summary['demand_weighted_skill_coverage']:.4f}",
        f"{summary['average_job_fit_score']:.4f}",
        success_count(summary, "job_fit_success_rate_060"),
        success_count(summary, "job_fit_success_rate_080"),
    )


def human_join(items: list[str]) -> str:
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + f", and {items[-1]}"


def summarize_seed_titles(seed_jobs: list[dict]) -> str:
    counts: Counter[str] = Counter()
    ordered_titles: list[str] = []
    for job in seed_jobs:
        title = job.get("title", "").strip()
        if not title:
            continue
        if title not in counts:
            ordered_titles.append(title)
        counts[title] += 1

    phrases: list[str] = []
    for title in ordered_titles:
        count = counts[title]
        if count == 1:
            phrases.append(f"`{title}`")
        elif count == 2:
            phrases.append(f"two `{title}` roles")
        else:
            phrases.append(f"{count} `{title}` roles")
    return human_join(phrases)


def summarize_success_examples(rows: list[dict], *, limit: int = 5) -> str:
    sorted_rows = sorted(rows, key=lambda row: row["average_job_fit_score"], reverse=True)
    examples: list[str] = []
    for row in sorted_rows:
        if row["average_job_fit_score"] < 0.60:
            continue
        company = str(row.get("company", "")).strip()
        title = str(row.get("title", "")).strip()
        if not company or company.lower() == "nan" or not title:
            continue
        examples.append(f"`{title}` at `{company}`")
        if len(examples) >= limit:
            break
    return human_join(examples)


def build_dynamic_content(payload: dict) -> dict:
    course = payload["course_level_results"]
    student = payload["student_case_results"]
    course_variants = course["variant_summaries"]
    student_variants = student["variant_summaries"]

    course_rows = [format_row(variant, summary) for variant, summary in course_variants.items()]
    student_rows = [format_row(variant, summary) for variant, summary in student_variants.items()]
    seed_title_summary = summarize_seed_titles(student["seed_jobs"])
    success_example_summary = summarize_success_examples(student["variant_results"]["KG + B_S + J_S"])

    course_best = course_variants["KG + J_S"]
    course_full = course_variants["KG + B_S + J_S"]
    student_best = student_variants["KG + B_S + J_S"]
    student_seed_only = student_variants["KG + J_S"]

    results_intro = (
        "This section reports two complementary results from the Big Data case study. "
        "It first evaluates course-level held-out alignment on a strict Big Data / Data Engineering benchmark, "
        "and then reports a student-centered opportunity-spillover analysis. Together, these results show both "
        "how curriculum enrichment improves alignment with external job requirements and how skills learned from "
        "a student's 7 selected target job postings can also transfer to other related jobs. "
        "The presentation is intentionally table-first so that the exact comparisons remain easy to read."
    )

    course_paragraphs = [
        (
            f"The course-level benchmark uses {len(payload['raw_jobs'])} deduplicated jobs in the broad pool, "
            f"filters them down to {len(payload['main_jobs'])} strict-scope Big Data / Data Engineering roles, "
            f"and evaluates {len(course['held_out_jobs'])} held-out jobs after a {len(course['build_jobs'])}-job build split. "
            f"In this held-out benchmark, the baseline KG variant achieved demand-weighted skill coverage of "
            f"{course_variants['KG']['demand_weighted_skill_coverage']:.4f} and an average job-fit score of "
            f"{course_variants['KG']['average_job_fit_score']:.4f}. Adding textbook skills raised these values to "
            f"{course_variants['KG + B_S']['demand_weighted_skill_coverage']:.4f} and "
            f"{course_variants['KG + B_S']['average_job_fit_score']:.4f}, respectively."
        ),
        (
            f"Market-skill enrichment produced the strongest held-out result. The KG + J_S variant reached "
            f"{course_best['demand_weighted_skill_coverage']:.4f} coverage and {course_best['average_job_fit_score']:.4f} "
            f"average job-fit, and it was the only course-level variant to produce a held-out success at the 0.80 threshold "
            f"({success_count(course_best, 'job_fit_success_rate_080')}). The full variant, KG + B_S + J_S, remained well above "
            f"the baseline at {course_full['demand_weighted_skill_coverage']:.4f} coverage and "
            f"{course_full['average_job_fit_score']:.4f} average job-fit, but it did not surpass KG + J_S on this particular split. "
            f"This pattern suggests that market-derived skills explain most of the measurable held-out lift, while textbook skills add "
            f"useful breadth without always improving every external benchmark slice."
        ),
    ]

    student_paragraphs = [
        (
            f"The personalized case study uses {len(student['seed_jobs'])} job postings explicitly selected by STUDENT(id=2) "
            f"as the seed portfolio and evaluates transfer on the remaining {len(student['remaining_jobs'])} jobs in the broad job graph. "
            f"Here the baseline KG variant achieved demand-weighted skill coverage of "
            f"{student_variants['KG']['demand_weighted_skill_coverage']:.4f} and an average job-fit score of "
            f"{student_variants['KG']['average_job_fit_score']:.4f}. Adding the student's selected book skills produced only a small increase "
            f"to {student_variants['KG + B_S']['demand_weighted_skill_coverage']:.4f} coverage and "
            f"{student_variants['KG + B_S']['average_job_fit_score']:.4f} average job-fit."
        ),
        (
            f"The dominant gain comes from the skills extracted from the student's 7 selected target job postings. "
            f"Using those seed-job skills in KG + J_S raises coverage to {student_seed_only['demand_weighted_skill_coverage']:.4f} "
            f"and average job-fit to {student_seed_only['average_job_fit_score']:.4f}, increasing the number of remaining jobs at "
            f">= 0.60 from {success_count(student_variants['KG'], 'job_fit_success_rate_060')} to "
            f"{success_count(student_seed_only, 'job_fit_success_rate_060')}. The full personalized variant, KG + B_S + J_S, achieves "
            f"the strongest overall student-case result at {student_best['demand_weighted_skill_coverage']:.4f} coverage and "
            f"{student_best['average_job_fit_score']:.4f} average job-fit. In simple terms, the skills learned for 7 target job postings "
            f"also make the student more suitable for many other related jobs because those jobs share overlapping competencies."
        ),
        (
            f"The seed portfolio itself spans {seed_title_summary}. That means the transferred skill bundle is not narrowly limited "
            f"to one job family: it mixes AI architecture and governance, SQL and data profiling, ETL and robust data pipelines, "
            f"Python/PySpark and Databricks, APIs and microservices, Linux or containerization, and testing or automation practices. "
            f"The spillover therefore appears in remaining roles such as {success_example_summary}, where overlapping competencies "
            f"like cross-functional data collaboration, relational database work, cloud-native engineering, and Python-based implementation "
            f"can be reused beyond the student's original targets."
        ),
    ]

    discussion_paragraphs = [
        (
            "Taken together, the two analyses show both system-level and student-level value. "
            "The course-level benchmark demonstrates that curriculum enrichment improves held-out alignment with relevant labor-market "
            "requirements, and the student-centered table shows that a focused target-job portfolio can unlock broader opportunity spillover. "
            "Across both analyses, market-derived job skills explain most of the measurable lift, while book skills contribute smaller but still complementary gains."
        ),
        (
            f"These findings should be interpreted with appropriate caution. The course-level benchmark remains a small held-out case study "
            f"with only {len(course['held_out_jobs'])} held-out jobs, and the student table is a personalized scenario analysis rather than a held-out benchmark. "
            "The evaluation measures estimated curriculum-to-market alignment rather than employment outcomes, so it is appropriate to claim improved visibility "
            "into skill gaps and partial evidence of transfer, but not employability gains, hiring success, or causal labor-market impact."
        ),
    ]

    return {
        "results_intro": results_intro,
        "course_rows": course_rows,
        "student_rows": student_rows,
        "course_paragraphs": course_paragraphs,
        "student_paragraphs": student_paragraphs,
        "discussion_paragraphs": discussion_paragraphs,
    }


def set_run_font(run, *, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def set_body_paragraph(paragraph) -> None:
    paragraph.style = "Body Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0.25)


def set_no_indent(paragraph, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0)
    fmt.right_indent = Inches(0)
    fmt.first_line_indent = Inches(0)
    if align is not None:
        paragraph.alignment = align


def insert_body_paragraph(anchor_paragraph, text: str):
    p = anchor_paragraph.insert_paragraph_before("", style="Body Text")
    set_body_paragraph(p)
    set_run_font(p.add_run(text))
    return p


def insert_heading(anchor_paragraph, text: str):
    p = anchor_paragraph.insert_paragraph_before(text, style="Heading 2")
    p.paragraph_format.keep_with_next = False
    return p


def insert_caption(anchor_paragraph, label: str, caption: str) -> None:
    p1 = anchor_paragraph.insert_paragraph_before("", style="Body Text")
    set_no_indent(p1, WD_ALIGN_PARAGRAPH.LEFT)
    p1.paragraph_format.keep_with_next = True
    set_run_font(p1.add_run(label), bold=True)

    p2 = anchor_paragraph.insert_paragraph_before("", style="Body Text")
    set_no_indent(p2, WD_ALIGN_PARAGRAPH.LEFT)
    p2.paragraph_format.keep_with_next = False
    set_run_font(p2.add_run(caption), italic=True)


def set_table_borders(table, *, outer: str = "single", inner_v: str | None = None, inner_h: str | None = None, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def set_edge(name: str, val: str) -> None:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")

    set_edge("top", outer)
    set_edge("bottom", outer)
    set_edge("left", outer if outer != "none" else "none")
    set_edge("right", outer if outer != "none" else "none")
    set_edge("insideH", inner_h or "none")
    set_edge("insideV", inner_v or "none")


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_cell_borders(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, val in kwargs.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), val or "none")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Body Text"
    p.alignment = align
    set_no_indent(p, align)
    set_run_font(p.add_run(text), size=font_size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_column_widths(table, widths: list[float]) -> None:
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
        for cell in table.columns[i].cells:
            cell.width = Inches(width)


def insert_table_before(anchor_paragraph, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    doc = anchor_paragraph.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, outer="none", inner_v="none", inner_h="none")
    set_table_width(table, sum(widths))

    set_column_widths(table, widths)

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_borders(table.rows[0].cells[i], top="single", bottom="single", left="none", right="none")

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.rows[row_idx].cells[col_idx], value, align=align, font_size=9.0)

    for cell in table.rows[-1].cells:
        set_cell_borders(cell, bottom="single", left="none", right="none")

    anchor_paragraph._element.addprevious(table._element)


def rewrite_body_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    set_body_paragraph(paragraph)
    set_run_font(paragraph.add_run(text))


def rewrite_paragraph_keep_style(paragraph, text: str) -> None:
    paragraph.clear()
    set_run_font(paragraph.add_run(text))


def startswith_any(text: str, prefixes: tuple[str, ...]) -> bool:
    return any(text.startswith(prefix) for prefix in prefixes)


def remove_between(start_paragraph, end_paragraph) -> None:
    body = start_paragraph._element.getparent()
    children = list(body)
    start_idx = children.index(start_paragraph._element)
    end_idx = children.index(end_paragraph._element)
    for child in children[start_idx + 1:end_idx]:
        body.remove(child)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def remove_results_figure_artifacts(doc: Document) -> None:
    removal_texts = {
        "Figure 2",
        "Course-level held-out market alignment across curriculum enrichment variants.",
        "Figure 3",
        "Course-level skill-gap closure across curriculum enrichment variants.",
    }
    results_heading_idx = next(
        index for index, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "4. Results"
    )
    student_heading_idx = next(
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip() == "4.2 Student-Centered Opportunity Spillover"
    )
    for paragraph in list(doc.paragraphs[results_heading_idx:student_heading_idx]):
        text = paragraph.text.strip()
        if text in removal_texts or paragraph_has_drawing(paragraph):
            paragraph._element.getparent().remove(paragraph._element)


def normalize_heading_layout(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 2":
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.page_break_before = False


def build_table_1_rows(payload: dict) -> list[tuple[str, str]]:
    return [
        ("Course", "Big Data"),
        ("Course chapters", str(COURSE_CHAPTER_COUNT)),
        ("Uploaded transcript documents", str(UPLOADED_TRANSCRIPT_DOCUMENTS)),
        ("Approved books", str(APPROVED_BOOKS)),
        ("Broad-pool deduplicated jobs", str(len(payload["raw_jobs"]))),
        ("Strict-scope Big Data jobs", str(len(payload["main_jobs"]))),
        ("Excluded broad-pool jobs", str(len(payload["main_dropped_jobs"]))),
        ("Build jobs", str(len(payload["build_jobs"]))),
        ("Held-out jobs", str(len(payload["held_out_jobs"]))),
        ("Student seed jobs", str(len(payload["student_seed_jobs"]))),
        ("Student remaining jobs", str(len(payload["student_remaining_jobs"]))),
    ]


def normalize_overview_table(doc: Document) -> None:
    target_headers = ["Task", "Section", "Agent / Component", "Core Purpose", "Depends On"]
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers != target_headers:
            continue
        set_table_borders(table, outer="none", inner_v="none", inner_h="none")
        set_table_width(table, 5.5)
        set_column_widths(table, [1.55, 0.7, 1.3, 2.1, 0.85])
        header_values = ["Task", "Sec.", "Agent /\nComponent", "Core Purpose", "Depends\nOn"]
        for i, value in enumerate(header_values):
            set_cell_text(table.rows[0].cells[i], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.25)
            set_cell_borders(table.rows[0].cells[i], top="single", bottom="single", left="none", right="none")
        for row in table.rows[1:]:
            for i, cell in enumerate(row.cells):
                set_cell_text(cell, cell.text.strip(), align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.75)
        for cell in table.rows[-1].cells:
            set_cell_borders(cell, bottom="single", left="none", right="none")
        break


def normalize_formula_tables(doc: Document) -> None:
    for table in doc.tables:
        if len(table.rows) != 1 or len(table.columns) != 2:
            continue
        right_text = table.rows[0].cells[1].text.strip()
        match = re.fullmatch(r"(?:Formula )?\((\d+)\)", right_text)
        if not match:
            continue
        set_table_borders(table, outer="none", inner_v="none", inner_h="none")
        set_table_width(table, 5.5)
        set_column_widths(table, [5.05, 0.45])
        left, right = table.rows[0].cells
        left_paragraph = left.paragraphs[0]
        left_paragraph.style = "Body Text"
        set_no_indent(left_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        left_paragraph.paragraph_format.keep_with_next = False
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_text(right, f"({match.group(1)})", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10.5)


def normalize_dataset_table(doc: Document, payload: dict) -> None:
    target_rows = build_table_1_rows(payload)
    for table in doc.tables:
        if len(table.columns) != 2 or table.rows[0].cells[0].text.strip() != "Field":
            continue
        if len(table.rows) != 1 + len(target_rows):
            continue
        set_table_borders(table, outer="none", inner_v="none", inner_h="none")
        set_table_width(table, 5.45)
        set_column_widths(table, [3.35, 2.1])
        set_cell_text(table.rows[0].cells[0], "Field", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)
        set_cell_text(table.rows[0].cells[1], "Value", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)
        set_cell_borders(table.rows[0].cells[0], top="single", bottom="single", left="none", right="none")
        set_cell_borders(table.rows[0].cells[1], top="single", bottom="single", left="none", right="none")
        for row, values in zip(table.rows[1:], target_rows):
            set_cell_text(row.cells[0], values[0], align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9.0)
            set_cell_text(row.cells[1], values[1], align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.0)
        for cell in table.rows[-1].cells:
            set_cell_borders(cell, bottom="single", left="none", right="none")
        break


def resize_inline_shape(shape, *, width_in: float) -> None:
    ratio = shape.height / shape.width
    shape.width = Inches(width_in)
    shape.height = int(shape.width * ratio)


def normalize_drawing_layout(doc: Document) -> None:
    drawing_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph_has_drawing(paragraph)]
    for paragraph in drawing_paragraphs:
        set_no_indent(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(6)


def normalize_inline_shapes(doc: Document) -> None:
    total_shapes = len(doc.inline_shapes)
    for index, shape in enumerate(doc.inline_shapes):
        if total_shapes >= 6 and index == 0:
            resize_inline_shape(shape, width_in=5.15)
            continue
        if total_shapes >= 6 and index == 1:
            resize_inline_shape(shape, width_in=5.55)
            continue
        if total_shapes >= 6 and index == total_shapes - 1:
            resize_inline_shape(shape, width_in=5.05)
            continue
        if total_shapes >= 5 and index == 0:
            resize_inline_shape(shape, width_in=5.55)
            continue
        if total_shapes >= 5 and index == total_shapes - 1:
            resize_inline_shape(shape, width_in=5.05)
            continue
        if total_shapes >= 5:
            resize_inline_shape(shape, width_in=4.95)
            continue
        resize_inline_shape(shape, width_in=5.5)


def normalize_references(doc: Document) -> None:
    references_start = None
    references_end = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "References":
            references_start = index + 1
            continue
        if references_start is not None and text == "Author short bio":
            references_end = index
            break

    if references_start is None:
        return

    end_index = references_end if references_end is not None else len(doc.paragraphs)
    for paragraph in doc.paragraphs[references_start:end_index]:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = paragraph.paragraph_format
        fmt.left_indent = Inches(0)
        fmt.right_indent = Inches(0)
        fmt.first_line_indent = Inches(0)
        fmt.keep_with_next = False
        for run in paragraph.runs:
            set_run_font(run)


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    payload = load_artifact()
    content = build_dynamic_content(payload)

    backup_path = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}.backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{DOCX_PATH.suffix}"
    )
    backup_path.write_bytes(DOCX_PATH.read_bytes())

    doc = Document(DOCX_PATH)
    results_heading = next(p for p in doc.paragraphs if p.text.strip() == "4. Results")
    conclusion_heading = next(p for p in doc.paragraphs if p.text.strip() == "5. Conclusion")

    remove_between(results_heading, conclusion_heading)

    insert_body_paragraph(conclusion_heading, content["results_intro"])

    insert_heading(conclusion_heading, "4.1 Course-Level Held-Out Alignment")
    for paragraph in content["course_paragraphs"]:
        insert_body_paragraph(conclusion_heading, paragraph)
    insert_caption(conclusion_heading, "Table 2", "Course-level held-out alignment results.")
    insert_table_before(
        conclusion_heading,
        ["Variant", "Coverage", "Average job-fit", ">= 0.60", ">= 0.80"],
        content["course_rows"],
        widths=[1.3, 1.05, 1.25, 0.95, 0.95],
    )

    insert_heading(conclusion_heading, "4.2 Student-Centered Opportunity Spillover")
    for paragraph in content["student_paragraphs"]:
        insert_body_paragraph(conclusion_heading, paragraph)
    insert_caption(conclusion_heading, "Table 3", "Student-centered opportunity-spillover results.")
    insert_table_before(
        conclusion_heading,
        ["Variant", "Coverage", "Average job-fit", ">= 0.60", ">= 0.80"],
        content["student_rows"],
        widths=[1.3, 1.05, 1.25, 0.95, 0.95],
    )

    insert_heading(conclusion_heading, "4.3 Discussion and Limitations")
    for paragraph in content["discussion_paragraphs"]:
        insert_body_paragraph(conclusion_heading, paragraph)

    conclusion_index = next(
        index for index, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "5. Conclusion"
    )
    for paragraph in doc.paragraphs[conclusion_index + 1:conclusion_index + 8]:
        if paragraph.text.strip().startswith(CONCLUSION_RESULTS_PREFIX):
            rewrite_body_paragraph(paragraph, CONCLUSION_RESULTS_PARAGRAPH)
            break

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if startswith_any(text, ABSTRACT_PREFIXES):
            rewrite_paragraph_keep_style(paragraph, CURRENT_ABSTRACT_PARAGRAPH)
        elif startswith_any(text, DATASET_SUMMARY_PREFIXES):
            rewrite_body_paragraph(paragraph, CURRENT_DATASET_SUMMARY_PARAGRAPH)
        elif startswith_any(text, EVALUATION_METHOD_PREFIXES):
            rewrite_body_paragraph(paragraph, CURRENT_EVALUATION_METHOD_PARAGRAPH)
        elif startswith_any(text, CASE_STUDY_QUESTION_PREFIXES):
            rewrite_body_paragraph(paragraph, CURRENT_CASE_STUDY_QUESTION_PARAGRAPH)

    normalize_heading_layout(doc)
    normalize_overview_table(doc)
    normalize_formula_tables(doc)
    normalize_dataset_table(doc, payload)
    normalize_drawing_layout(doc)
    normalize_inline_shapes(doc)
    normalize_references(doc)
    remove_results_figure_artifacts(doc)
    doc.save(DOCX_PATH)
    print(f"Updated {DOCX_PATH}")
    print(f"Backup written to {backup_path}")


if __name__ == "__main__":
    main()
