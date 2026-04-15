from __future__ import annotations

import re
import shutil
import struct
import subprocess
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


ROOT = Path("/Users/khajievroma/Projects/lab_tutor/docs/conference_paper")
TEMPLATE = ROOT / "教育技术协会 英文论文模板 2026.docx"
OUTPUT = ROOT / "lab_tutor_conference_paper_working.docx"
MARKDOWN = ROOT / "lab_tutor_conference_paper_final.md"
ALGORITHM_RENDERER = ROOT / "render_algorithm_assets.py"
GENERATED_ALGORITHM_DIR = ROOT / "generated" / "algorithms"
COMMON_BIN_DIRS = (
    Path("/opt/homebrew/bin"),
    Path("/usr/local/bin"),
    Path("/Library/TeX/texbin"),
)

FIGURES = {
    1: {
        "path": ROOT / "system-architecture.png",
        "caption": "Task-level multi-agent system overview showing the four specialized agents, their sequential dependencies, and their shared updates to the unified Neo4j curriculum knowledge graph.",
        "width": 5.55,
    },
    2: {
        "path": ROOT / "logic_architecture.png",
        "caption": "Detailed agent-level architecture showing the four specialized agents, their key workflow steps, and their data flows into and out of the shared Neo4j curriculum knowledge graph.",
        "width": 5.15,
    },
    3: {
        "path": ROOT / "UI_schreenshot_of_student_choosing_skills.png",
        "caption": "Student-facing skill selection and learning resource interface, showing graph-linked skill choices and recommended readings derived from the enriched curriculum knowledge graph.",
        "width": 5.05,
    },
}

ALGORITHMS = {
    1: {
        "caption": "Curricular Alignment Architect: Book-Skill Bank Construction",
        "lines": [
            "Require: teacher-approved books B_sel, course chapter scaffold C, course subject sigma.",
            "Ensure: book skill bank S_B mapped into course chapters.",
            "1. For each approved book B_i, load extracted chapters ch_1 ... ch_N.",
            "2. For each unfinished chapter ch_j in parallel, extract chapter summary, practical skills, and prerequisite concepts.",
            "3. Judge the extraction against chapter evidence and run at most one revision pass if needed.",
            "4. Embed skill names and descriptions, then save chapter outputs to PostgreSQL and Neo4j.",
            "5. Create BOOK_SKILL and REQUIRES_CONCEPT relations for accepted chapter outputs.",
            "6. Load the teacher-owned course chapter scaffold and all extracted book chapters with skills.",
            "7. For each book chapter, map each extracted book skill to the best-fit course chapter or chapters.",
            "8. Write BOOK_SKILL -> MAPPED_TO -> COURSE_CHAPTER relations.",
            "9. Return the mapped book skill bank S_B.",
        ],
    },
    2: {
        "caption": "Market Demand Analyst: Market-Skill Bank Construction",
        "lines": [
            "Require: course context D, approved search terms T, curriculum graph G_course.",
            "Ensure: cleaned market skill bank S_M linked to chapters, jobs, and concepts.",
            "1. Fetch job postings for approved search terms across supported job sites in parallel.",
            "2. Deduplicate postings by normalized title and company, then group them for teacher review.",
            "3. Keep only job groups judged relevant to the target course domain.",
            "4. For each retained job, extract competency-style market skills from the job description.",
            "5. Aggregate, frequency-score, and canonicalize the extracted market-skill set.",
            "6. If the candidate set is large enough, run an LLM-based merge pass and embedding-based deduplication.",
            "7. Let the teacher curate the cleaned skill list by name, category, and priority.",
            "8. For each curated skill, check graph coverage and map it to the best-fit course chapter.",
            "9. Remove redundant skills relative to existing curriculum coverage.",
            "10. Link retained skills to prerequisite concepts and source jobs, then insert MARKET_SKILL nodes into Neo4j.",
            "11. Return the final market skill bank S_M.",
        ],
    },
    3: {
        "caption": "Skill-Conditioned Reading and Video Retrieval",
        "lines": [
            "Require: selected skills S_sel, skill profiles P, curriculum graph G.",
            "Ensure: reading and video resources linked to skills.",
        ],
    },
}

TABLE_1_ROWS = [
    ("Course", "Big Data"),
    ("Course chapters", "6"),
    ("Uploaded transcript documents", "38"),
    ("Approved books", "3"),
    ("Deduplicated broad-pool jobs", "62"),
    ("Strict-scope Big Data jobs", "28"),
    ("Build jobs", "20"),
    ("Held-out jobs", "8"),
    ("Student seed jobs", "7"),
    ("Student remaining jobs", "55"),
    ("Extracted course skills", "256"),
]

TABLE_2_ROWS = [
    ("KG", "0.2961", "0.2775", "0/8", "0/8"),
    ("KG + B_S", "0.3398", "0.3179", "1/8", "0/8"),
    ("KG + J_S", "0.4806", "0.4602", "1/8", "1/8"),
    ("KG + B_S + J_S", "0.4757", "0.4568", "1/8", "0/8"),
]

TABLE_3_ROWS = [
    ("KG", "0.2608", "0.2588", "1/55", "0/55"),
    ("KG + B_S", "0.2671", "0.2646", "1/55", "0/55"),
    ("KG + J_S", "0.4196", "0.4259", "6/55", "0/55"),
    ("KG + B_S + J_S", "0.4266", "0.4333", "6/55", "0/55"),
]

TASK_MAPPING_ROWS = [
    ("Task 1", "Document extraction: teacher materials -> course chapters, concepts, embeddings", "Concept/chapter review"),
    ("Task 2", "Curricular Alignment Architect: approved books -> mapped BOOK_SKILL layer", "Textbook approval"),
    ("Task 3", "Market Demand Analyst: filtered job postings -> mapped MARKET_SKILL layer", "Job relevance and skill curation"),
    ("Task 4", "Textual Resource Analyst + Video Agent: enriched skill set -> readings, videos, questions", "Optional resource review"),
]

FORMULA_LATEX = {
    1: r"S_{book}=0.30C_{topic}+0.20C_{struc}+0.15C_{scope}+0.15C_{pub}+0.10C_{auth}+0.10C_{time}",
    2: r"S_{res}=0.15C_{rec}+0.25C_{cov}+0.15C_{emb}+0.20C_{ped}+0.15C_{depth}+0.10C_{extra}",
    3: r"Coverage(v)=\frac{\sum_s w_s\, y(v,s)}{\sum_s w_s}",
    4: r"JobFit_j(v)=\frac{1}{|S_j|}\sum_{s \in S_j} y(v,s)",
    5: r"AvgJobFit(v)=\frac{1}{|J|}\sum_{j \in J} JobFit_j(v)",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def resolve_command(name: str) -> str:
    resolved = shutil.which(name)
    if resolved:
        return resolved
    for directory in COMMON_BIN_DIRS:
        candidate = directory / name
        if candidate.exists():
            return str(candidate)
    raise FileNotFoundError(f"Required command not found: {name}")


def delete_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def clear_paragraph(paragraph) -> None:
    paragraph.text = ""


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    def replace_link(match: re.Match[str]) -> str:
        label, url = match.group(1), match.group(2)
        return url if label == url else f"{label} ({url})"

    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", replace_link, text)
    return " ".join(text.split())


def add_inline_markup(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def set_body_paragraph(paragraph) -> None:
    paragraph.style = "Body Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0.25)


def set_no_indent(paragraph, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0)
    fmt.right_indent = Inches(0)
    fmt.first_line_indent = Inches(0)
    if align is not None:
        paragraph.alignment = align


def add_body_text(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_body_paragraph(p)
    p.add_run(strip_inline_markup(text))


def add_heading(doc: Document, text: str, level: int) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = level == 1
    p.add_run(text)


def set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.style = "Body Text"
    set_no_indent(p, align)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, *, outer: str = "single", inner_v: str | None = None, inner_h: str | None = None, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def set_edge(name: str, val: str) -> None:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")

    set_edge("top", outer)
    set_edge("bottom", outer)
    set_edge("left", outer if outer != "none" else "none")
    set_edge("right", outer if outer != "none" else "none")
    set_edge("insideH", inner_h or "none")
    set_edge("insideV", inner_v or "none")


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_cell_borders(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, val in kwargs.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        if val is None:
            val = "none"
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")


def add_caption(doc: Document, label: str, caption: str) -> None:
    p1 = doc.add_paragraph(style="Body Text")
    set_no_indent(p1, WD_ALIGN_PARAGRAPH.LEFT)
    p1.paragraph_format.keep_with_next = True
    r1 = p1.add_run(label)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r1.font.size = Pt(10.5)

    p2 = doc.add_paragraph(style="Body Text")
    set_no_indent(p2, WD_ALIGN_PARAGRAPH.LEFT)
    p2.paragraph_format.keep_with_next = False
    r2 = p2.add_run(caption)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r2.font.size = Pt(10.5)


def add_minimal_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float], note_small: bool = False) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, outer="none", inner_v="none", inner_h="none")
    set_table_width(table, sum(widths))

    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)
        set_cell_borders(table.rows[0].cells[i], top="single", bottom="single", left="none", right="none")

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, len(row_values) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            size = 8.5 if note_small and col_idx == len(row_values) - 1 else 9
            set_cell_text(table.rows[row_idx].cells[col_idx], value, align=align, font_size=size)

    last_row = table.rows[-1]
    for cell in last_row.cells:
        set_cell_borders(cell, bottom="single", left="none", right="none")


def add_task_mapping_table(doc: Document) -> None:
    add_minimal_table(
        doc,
        ["Task", "Flow", "Review"],
        TASK_MAPPING_ROWS,
        widths=[0.85, 3.35, 1.3],
        note_small=True,
    )


def add_figure(doc: Document, number: int, image_path: Path, caption: str, width_inches: float) -> None:
    p = doc.add_paragraph(style="Body Text")
    set_no_indent(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(fit_image_width(image_path, max_width=width_inches, max_height=4.7)))
    add_caption(doc, f"Figure {number}", caption)


def add_algorithm(doc: Document, number: int, caption: str, lines: list[str]) -> None:
    image_path = ensure_algorithm_asset(number)
    p = doc.add_paragraph(style="Body Text")
    set_no_indent(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(fit_image_width(image_path, max_width=5.0, max_height=4.8)))


def add_formula(doc: Document, formula: str, number: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, outer="none", inner_v="none", inner_h="none")
    set_table_width(table, 5.5)
    left, right = table.rows[0].cells
    table.columns[0].width = Inches(5.05)
    table.columns[1].width = Inches(0.45)
    left.width = Inches(5.05)
    right.width = Inches(0.45)

    left.text = ""
    p = left.paragraphs[0]
    p.style = "Body Text"
    set_no_indent(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = True
    p._p.append(latex_to_omml_paragraph(FORMULA_LATEX.get(number, formula)))

    set_cell_text(right, f"({number})", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10.5)


def fit_image_width(image_path: Path, *, max_width: float, max_height: float) -> float:
    width_px, height_px = read_image_size(image_path)
    if not width_px or not height_px:
        return max_width
    aspect_ratio = width_px / height_px
    return min(max_width, max_height * aspect_ratio)


def read_image_size(image_path: Path) -> tuple[int, int]:
    suffix = image_path.suffix.lower()
    if suffix == ".png":
        with image_path.open("rb") as f:
            header = f.read(24)
        if header[:8] != b"\x89PNG\r\n\x1a\n":
            raise ValueError(f"Invalid PNG header for {image_path}")
        width, height = struct.unpack(">II", header[16:24])
        return width, height
    if suffix in {".jpg", ".jpeg"}:
        with image_path.open("rb") as f:
            data = f.read()
        index = 2
        while index < len(data):
            if data[index] != 0xFF:
                index += 1
                continue
            marker = data[index + 1]
            if 0xC0 <= marker <= 0xC3:
                height, width = struct.unpack(">HH", data[index + 5:index + 9])
                return width, height
            block_length = struct.unpack(">H", data[index + 2:index + 4])[0]
            index += 2 + block_length
    raise ValueError(f"Unsupported image type for size detection: {image_path}")


def latex_to_omml_paragraph(latex: str):
    md = f"$$\n{latex}\n$$\n"
    with TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        md_path = tmpdir / "equation.md"
        docx_path = tmpdir / "equation.docx"
        md_path.write_text(md)
        subprocess.run(
            [resolve_command("pandoc"), str(md_path), "-o", str(docx_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        with ZipFile(docx_path) as zf:
            xml = zf.read("word/document.xml")

    root = etree.fromstring(xml)
    ns = {
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    omml = root.xpath("//m:oMathPara", namespaces=ns)[0]
    return etree.fromstring(etree.tostring(omml))


def ensure_algorithm_asset(number: int) -> Path:
    output_png = GENERATED_ALGORITHM_DIR / f"algorithm_{number}.png"
    if output_png.exists():
        return output_png
    subprocess.run(
        [
            resolve_command("python3"),
            str(ALGORITHM_RENDERER),
            "--algorithms",
            str(number),
            "--width-in",
            "7.0",
            "--dpi",
            "600",
            "--engine",
            "auto",
        ],
        cwd=ROOT,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    if not output_png.exists():
        raise FileNotFoundError(f"Rendered algorithm asset not found: {output_png}")
    return output_png


def parse_markdown_body(md_text: str):
    before_refs, refs_text = md_text.split("\n## References\n", 1)
    lines = before_refs.splitlines()

    title = lines[0].lstrip("# ").strip()
    abstract = ""
    keywords = ""

    body_start = 0
    for i, line in enumerate(lines):
        if line.strip() == "## Abstract":
            abstract = lines[i + 2].strip()
        if line.strip() == "## Keywords":
            keywords = lines[i + 2].strip()
        if line.strip().startswith("## 1. Introduction"):
            body_start = i
            break

    body_lines = lines[body_start:]
    blocks: list[tuple[str, str]] = []
    current: list[str] = []

    def flush_para():
        if current:
            blocks.append(("para", strip_inline_markup(" ".join(current).strip())))
            current.clear()

    for line in body_lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            flush_para()
            blocks.append(("h1", stripped[3:].strip()))
        elif stripped.startswith("### "):
            flush_para()
            blocks.append(("h2", stripped[4:].strip()))
        elif re.match(r"^\[Figure \d+ about here: .+\]$", stripped):
            flush_para()
            blocks.append(("figure", stripped))
        elif re.match(r"^\[Table \d+ about here: .+\]$", stripped):
            flush_para()
            blocks.append(("table", stripped))
        elif re.match(r"^\[Algorithm \d+ about here: .+\]$", stripped):
            flush_para()
            blocks.append(("algorithm", stripped))
        elif re.match(r"^\[Task Mapping about here: .+\]$", stripped):
            flush_para()
            blocks.append(("task_mapping", stripped))
        elif stripped.startswith("`") and stripped.endswith("`"):
            flush_para()
            blocks.append(("formula", stripped.strip("`")))
        elif not stripped:
            flush_para()
        else:
            current.append(stripped)
    flush_para()

    references = [line.strip() for line in refs_text.splitlines() if line.strip()]
    return title, abstract, keywords, blocks, references


def populate_front_matter(doc: Document, title: str, abstract: str, keywords: str) -> None:
    doc.paragraphs[0].text = title
    doc.paragraphs[1].text = ""
    doc.paragraphs[2].text = "First AUTHOR1, Second AUTHOR2*, Third AUTHOR3"
    doc.paragraphs[3].text = "(1. affiliation, city, province, postcode; 2*. affiliation, city, province, postcode; 3. affiliation, city, province, postcode)"
    doc.paragraphs[4].text = "* Corresponding author"
    doc.paragraphs[5].text = f"Abstract: {abstract}"
    doc.paragraphs[6].text = f"Keywords: {keywords}"


def cleanup_template_tail(doc: Document) -> None:
    while len(doc.paragraphs) > 7:
        delete_paragraph(doc.paragraphs[-1])
    while doc.tables:
        delete_table(doc.tables[-1])


def add_table_block(doc: Document, number: int) -> None:
    if number == 1:
        add_caption(doc, "Table 1", "Big Data case study dataset summary.")
        add_minimal_table(
            doc,
            ["Field", "Value"],
            TABLE_1_ROWS,
            widths=[3.35, 2.0],
        )
    elif number == 2:
        add_caption(doc, "Table 2", "Course-level held-out alignment results.")
        add_minimal_table(
            doc,
            ["Variant", "Coverage", "Average job-fit", ">= 0.60", ">= 0.80"],
            TABLE_2_ROWS,
            widths=[1.3, 1.05, 1.25, 0.95, 0.95],
        )
    elif number == 3:
        add_caption(doc, "Table 3", "Student-centered opportunity-spillover results.")
        add_minimal_table(
            doc,
            ["Variant", "Coverage", "Average job-fit", ">= 0.60", ">= 0.80"],
            TABLE_3_ROWS,
            widths=[1.3, 1.05, 1.25, 0.95, 0.95],
        )


def add_references(doc: Document, references: list[str]) -> None:
    heading = doc.add_paragraph(style="中文参考文献")
    heading.text = "References"

    for ref in references:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0)
        fmt.first_line_indent = Inches(0)
        add_inline_markup(p, ref)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(10.5)


def add_back_matter(doc: Document) -> None:
    ack_h = doc.add_paragraph(style="Body Text")
    set_no_indent(ack_h, WD_ALIGN_PARAGRAPH.LEFT)
    ack_h.paragraph_format.keep_with_next = True
    run = ack_h.add_run("Acknowledgement")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    ack_p = doc.add_paragraph(style="Body Text")
    set_body_paragraph(ack_p)
    ack_p.add_run("Funding information or acknowledgement text will be added here if applicable.")

    doc.add_paragraph(style="Body Text")

    add_references(doc, parse_markdown_body(MARKDOWN.read_text())[4])

    bio_h = doc.add_paragraph(style="Body Text")
    set_no_indent(bio_h, WD_ALIGN_PARAGRAPH.LEFT)
    bio_h.paragraph_format.keep_with_next = True
    bio_run = bio_h.add_run("Author short bio")
    bio_run.bold = True
    bio_run.font.size = Pt(14)
    bio_run.font.name = "Times New Roman"
    bio_run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    bio_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    bio_p = doc.add_paragraph(style="Body Text")
    set_body_paragraph(bio_p)
    bio_p.add_run("Author biographies will be added after the author list and affiliations are finalized.")


def main() -> None:
    doc = Document(TEMPLATE)
    title, abstract, keywords, blocks, references = parse_markdown_body(MARKDOWN.read_text())
    populate_front_matter(doc, title, abstract, keywords)
    cleanup_template_tail(doc)

    formula_counter = 1

    for kind, payload in blocks:
        if kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "para":
            add_body_text(doc, payload)
        elif kind == "figure":
            num = int(re.search(r"\[Figure (\d+)", payload).group(1))
            fig = FIGURES[num]
            add_figure(doc, num, fig["path"], fig["caption"], fig["width"])
        elif kind == "table":
            num = int(re.search(r"\[Table (\d+)", payload).group(1))
            add_table_block(doc, num)
        elif kind == "algorithm":
            num = int(re.search(r"\[Algorithm (\d+)", payload).group(1))
            alg = ALGORITHMS[num]
            add_algorithm(doc, num, alg["caption"], alg["lines"])
        elif kind == "task_mapping":
            add_task_mapping_table(doc)
        elif kind == "formula":
            add_formula(doc, payload, formula_counter)
            formula_counter += 1

    add_back_matter(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
