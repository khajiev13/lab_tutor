#!/usr/bin/env python3
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/khajievroma/Projects/lab_tutor/docs/conference_paper")
SOURCE = Path(
    "/Users/khajievroma/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_803eyynpjdoh12_8abb/msg/file/2026-03/Draft_Roma-v3(2).docx"
)
OUTPUT = ROOT / "Draft_Roma-v4_teacher_feedback.docx"
FIGURE_SCRIPT = ROOT / "render_task_overview_figure.py"
FIGURE_PATH = ROOT / "system-architecture.png"

DATASET_ROWS = [
    ("Course", "Big Data"),
    ("Course chapters", "6"),
    ("Uploaded transcript documents", "38"),
    ("Approved books", "3"),
    ("Deduplicated broad-pool jobs", "62"),
    ("Strict-scope Big Data jobs", "28"),
    ("Build jobs", "20"),
    ("Held-out jobs", "8"),
    ("Student seed jobs", "7"),
    ("Student remaining jobs", "55"),
]

TASK_ROWS = [
    ("Task 1", "Document extraction: teacher materials -> course chapters, concepts, embeddings", "Concept/chapter review"),
    ("Task 2", "Curricular Alignment Architect: approved books -> mapped BOOK_SKILL layer", "Textbook approval"),
    ("Task 3", "Market Demand Analyst: filtered job postings -> mapped MARKET_SKILL layer", "Job relevance and skill curation"),
    ("Task 4", "Textual Resource Analyst + Video Agent: enriched skill set -> readings, videos, questions", "Optional resource review"),
]


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, *, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def rewrite_paragraph(
    paragraph: Paragraph,
    text: str,
    *,
    style: str = "Body Text",
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    size: float = 10.5,
    bold: bool = False,
) -> Paragraph:
    paragraph.text = text
    paragraph.style = style
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)
    return paragraph


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    *,
    style: str = "Body Text",
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    size: float = 10.5,
    bold: bool = False,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    rewrite_paragraph(new_para, text, style=style, alignment=alignment, size=size, bold=bold)
    return new_para


def move_table_after(table: Table, paragraph: Paragraph) -> None:
    paragraph._p.addnext(table._tbl)


def format_cell(cell, text: str, *, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, size: float = 9.5, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Body Text"
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    widths: list[float],
    *,
    small_last_column: bool = False,
) -> Table:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = doc.tables[5].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)

    for idx, header in enumerate(headers):
        format_cell(table.rows[0].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, bold=True)

    for row_idx, values in enumerate(rows, start=1):
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            size = 8.8 if small_last_column and col_idx == len(values) - 1 else 9.0
            format_cell(table.rows[row_idx].cells[col_idx], value, align=align, size=size)
    return table


def replace_embedded_figure(docx_path: Path, image_name: str, replacement_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    replacement = replacement_path.read_bytes()
    with ZipFile(docx_path) as src, ZipFile(temp_path, "w") as dst:
        for item in src.infolist():
            data = replacement if item.filename == image_name else src.read(item.filename)
            dst.writestr(item, data)
    temp_path.replace(docx_path)


def main() -> None:
    subprocess.run(["python3", str(FIGURE_SCRIPT)], cwd=ROOT, check=True)
    if OUTPUT.exists():
        OUTPUT.chmod(0o644)
        OUTPUT.unlink()
    shutil.copy2(SOURCE, OUTPUT)
    OUTPUT.chmod(0o644)

    doc = Document(OUTPUT)

    paragraphs = doc.paragraphs
    figure_caption = paragraphs[26]
    paragraph_3_8 = paragraphs[50]
    discussion_heading = paragraphs[65]

    rewrite_paragraph(
        paragraphs[19],
        "Lab Tutor is organized as a four-task multi-agent architecture because the curriculum problem is layered: teacher materials must first be structured, external textbook and market evidence must then be aligned, and only after that can the system generate learning support in a traceable way.",
    )
    rewrite_paragraph(
        paragraphs[20],
        "Figure 1 shows the logical architecture of the whole system. It makes the value of the design explicit: Task 1 builds the teacher-owned scaffold, Tasks 2 and 3 enrich that scaffold from textbooks and labor-market evidence, and Task 4 uses the enriched graph to support readings and videos.",
    )
    for idx in (21, 22, 23, 24):
        delete_paragraph(paragraphs[idx])
    rewrite_paragraph(figure_caption, "Figure 1. Task-level multi-agent system overview", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    task_intro = insert_paragraph_after(
        figure_caption,
        "The whole framework can be summarized as four coordinated tasks as follows.",
        style="Body Text",
    )
    task_table = build_table(
        doc,
        ["Task", "Flow", "Review"],
        TASK_ROWS,
        widths=[0.9, 3.35, 1.25],
        small_last_column=True,
    )
    move_table_after(task_table, task_intro)

    rewrite_paragraph(
        paragraphs[44],
        "Lab Tutor is open-source (https://github.com/khajiev13/lab_tutor; live at https://labtutor.app). Codebase: FastAPI backend, React frontend, Neo4j graph database. Agent workflows use LangGraph on Azure. Figure 2 makes the internal orchestration visible by showing how the four specialized agents read from and write to the shared Neo4j graph, while Figure 3 shows the student-facing interface that turns Task 4 resource curation into concrete learning support.",
    )

    rewrite_paragraph(
        paragraph_3_8,
        "We evaluate on a Big Data course. Four variants: KG (baseline), KG+B_S (textbook), KG+J_S (market), KG+B_S+J_S (full). Table 1 summarizes the case-study dataset and split configuration: 62 deduplicated jobs, 28 strict-scope Big Data roles, a 20/8 build-held-out split, and a student analysis with 7 seed jobs and 55 remaining jobs.",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    dataset_caption = insert_paragraph_after(
        paragraph_3_8,
        "Table 1 Big Data case study dataset summary.",
        style="Body Text",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    dataset_table = build_table(doc, ["Field", "Value"], DATASET_ROWS, widths=[3.25, 2.05])
    move_table_after(dataset_table, dataset_caption)

    rewrite_paragraph(paragraphs[58], "Table 2 summarizes results on 8 held-out Big Data jobs.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    rewrite_paragraph(paragraphs[59], "Table 2 Course-level held-out alignment results.", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    rewrite_paragraph(
        paragraphs[60],
        "Market-skill enrichment produced the strongest result (KG+J_S: 0.4806 coverage, 0.4602 job-fit). The full variant remained above baseline at 0.4757 coverage and 0.4568 job-fit, but it did not surpass KG+J_S on this split. A plausible explanation is that the held-out benchmark contains only 8 jobs and KG+J_S already injects the signal most directly matched to those jobs. In this setting, textbook skills still add chapter-level breadth and coherence, but some of that added coverage is overlapping or less hold-out-specific, so the combined variant can dip slightly without weakening the value of the textbook layer.",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    rewrite_paragraph(paragraphs[62], "Table 3 shows transfer from 7 student-selected jobs to 55 remaining jobs.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    rewrite_paragraph(paragraphs[63], "Table 3 Student-centered opportunity-spillover results.", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    rewrite_paragraph(
        paragraphs[64],
        "Skills from 7 target jobs improved average job-fit on the other 55 jobs from 0.2588 to 0.4259. The personalized book layer adds only a small gain from KG to KG+B_S because the student's 25 selected book skills are concentrated in concept-heavy Big Data, storage, and analysis chapters. By contrast, the same student's 38 selected market skills are much more skewed toward engineering practice - architecture, cloud, databases, APIs, pipelines, and implementation - so the stronger transfer signal comes from the target-job layer. This is why KG+J_S drives most of the lift and KG+B_S+J_S adds only a modest final increment.",
        style="Body Text",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    rewrite_paragraph(
        paragraphs[66],
        "Market-derived skills explain most measurable lift, while textbook skills contribute coherence and chapter structure. Taken together, the two tables suggest that textbook enrichment mainly strengthens pedagogical organization, whereas market and seed-job enrichment contribute the benchmark-facing engineering signal. Limitations: small held-out set (8 jobs), student table is a personalized scenario analysis, and the evaluation measures estimated alignment rather than employment outcomes.",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    doc.save(OUTPUT)
    replace_embedded_figure(OUTPUT, "word/media/image1.png", FIGURE_PATH)


if __name__ == "__main__":
    main()
