from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/khajievroma/Projects/lab_tutor/docs/conference_paper")
REVISION_DIR = ROOT / "revisions"
SOURCE_DOCX = Path(
    "/Users/khajievroma/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_803eyynpjdoh12_8abb/msg/file/2026-03/Draft_Roma-v3(2).docx"
)
OUTPUT_DOCX = REVISION_DIR / "Draft_Roma-v3_teacher_feedback_revised.docx"
FIGURE_PATH = REVISION_DIR / "figure1_task_overview.png"

FIGURE_CAPTION = (
    "Figure 1. Task-level overview of the Lab Tutor framework and dependencies "
    "among the four tasks."
)
TABLE_1_CAPTION = "Table 1. Task-to-figure mapping of the Lab Tutor framework."
COURSE_TABLE_LEAD = "Table 2 summarizes results on 8 held-out Big Data jobs."
COURSE_TABLE_CAPTION = "Table 2 Course-level held-out alignment results."
STUDENT_TABLE_LEAD = (
    "Table 3 shows transfer from 7 student-selected jobs to 55 remaining jobs."
)
STUDENT_TABLE_CAPTION = "Table 3 Student-centered opportunity-spillover results."

FRAMEWORK_PARAGRAPHS = [
    (
        "Figure 1 shows the logical architecture of Lab Tutor. The framework is "
        "organized around a shared Neo4j course knowledge graph that serves as the "
        "coordination layer for four specialized tasks. We designed the system in this "
        "way because course materials, textbook evidence, market skills, and learning "
        "resources should remain connected through explicit chapter-, concept-, and "
        "skill-level relations rather than isolated files or disconnected agent outputs."
    ),
    (
        "Based on Figure 1, the overall workflow can be summarized as four tasks. "
        "Task 1 constructs the transcript-grounded curriculum scaffold. Task 2 adds "
        "textbook-derived skills and prerequisite structure. Task 3 adds "
        "market-demanded skills from job postings. Task 4 uses the enriched skill "
        "graph to support downstream readings and videos. Table 1 makes the "
        "correspondence between each task, its agent responsibility, and its role in "
        "the framework explicit."
    ),
    (
        "The dependency structure in Figure 1 is also methodologically important. "
        "Task 1 is the prerequisite because later agents need a stable teacher-owned "
        "chapter scaffold. Tasks 2 and 3 are parallel enrichment paths that both "
        "write back to the same graph, making their contributions separable for "
        "evaluation. Task 4 is positioned last because it consumes the enriched graph "
        "produced by the earlier tasks. Human review checkpoints remain active at "
        "textbook approval, skill curation, and concept review so that curricular "
        "authority stays with the instructor."
    ),
]

COURSE_DISCUSSION = (
    "Market-skill enrichment produced the strongest held-out result (KG+J_S: 0.4806 "
    "coverage, 0.4602 job-fit). The slight non-monotonicity relative to KG+B_S+J_S is "
    "plausible in this split because the held-out set is very small (8 jobs) and is "
    "dominated by engineering-oriented roles such as Data Engineer, DevSecOps "
    "Engineer, and Senior Data Ops Engineer. In this setting, job-derived skills are "
    "more directly calibrated to the benchmark, whereas book-derived skills add "
    "broader chapter-structured coverage and some overlapping or more academic "
    "vocabulary. The full variant therefore remains essentially tied with KG+J_S, but "
    "does not exceed it on this specific hold-out set."
)

STUDENT_DISCUSSION = (
    "Skills from 7 target jobs (AI Development Architect, Data Architect, Data "
    "Engineer, Full Stack Engineer, Engineering Intern, and Junior Developer) improved "
    "average job-fit on other jobs from 0.2588 to 0.4259. The negligible gain from "
    "KG+B_S over KG suggests that the selected jobs lean more strongly toward "
    "engineering practice than toward textbook-style theory. In Neo4j, the student's "
    "selected book skills are concentrated in chapter-structured topics such as CAP "
    "theorem, clustering, visualization, MongoDB, and Cypher, whereas the selected "
    "market skills emphasize SQL, pipelines, cloud infrastructure, microservices, "
    "Linux, and CI/CD. This helps explain why textbook skills add only a small lift in "
    "the student spillover analysis: they strengthen conceptual coherence, but they "
    "are less directly aligned with the practice-oriented benchmark defined by the "
    "seven chosen jobs."
)

DISCUSSION_LIMITATIONS = (
    "Market-derived skills explain most measurable lift, while textbook-derived skills "
    "provide complementary pedagogical coverage whose measurable effect depends on the "
    "benchmark. In the course-level table, the full model is essentially tied with "
    "KG+J_S, which is consistent with overlap between book-side and job-side skills in "
    "a very small held-out split. In the student-centered table, textbook skills "
    "remain useful for conceptual depth, but the direct spillover gain comes mainly "
    "from the practice-oriented market skills selected from the student's target jobs. "
    "Limitations remain important: the held-out set contains only 8 jobs, and the "
    "evaluation measures estimated alignment rather than employment outcomes."
)

TASK_TABLE_ROWS = [
    (
        "Task 1",
        "Document extraction pipeline",
        "Parse teacher materials into course chapters, concepts, and embeddings",
        "Foundation for Tasks 2-4",
    ),
    (
        "Task 2",
        "Curricular Alignment Architect",
        "Discover books and map book skills plus prerequisites to chapters",
        "Parallel enrichment path from Task 1",
    ),
    (
        "Task 3",
        "Market Demand Analyst",
        "Extract and align market skills from job postings to chapters",
        "Parallel enrichment path from Task 1",
    ),
    (
        "Task 4",
        "Textual Resource Analyst + Video Agent",
        "Retrieve and rank readings and videos for selected skills",
        "Consumes the enriched graph",
    ),
]


def generate_figure(path: Path) -> None:
    dot = shutil.which("dot")
    if dot is None:
        raise FileNotFoundError("Graphviz 'dot' is required to generate Figure 1")

    dot_source = """
digraph LabTutor {
  graph [
    rankdir=TB,
    splines=ortho,
    bgcolor="white",
    pad=0.35,
    nodesep=0.45,
    ranksep=0.65,
    fontname="Helvetica"
  ];
  node [
    shape=box,
    style="rounded,filled",
    fontname="Helvetica",
    fontsize=22,
    margin="0.2,0.16",
    penwidth=2.2
  ];
  edge [
    color="#64748b",
    penwidth=2.0,
    arrowsize=0.8
  ];

  labtitle [
    shape=plain,
    label=<
      <FONT FACE="Helvetica" POINT-SIZE="28"><B>Lab Tutor Task-Level Framework Overview</B></FONT>
    >
  ];

  subgraph cluster_foundation {
    label="Knowledge foundation";
    style="rounded,filled";
    color="#d4d4d8";
    fillcolor="#f8fafc";
    fontsize=20;
    t1 [
      fillcolor="#dbeafe",
      color="#2563eb",
      label=<
        <B>Task 1</B><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="18">Document extraction</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Teacher materials -&gt; chapters, concepts, embeddings</FONT>
      >
    ];
  }

  subgraph cluster_alignment {
    label="Knowledge alignment";
    style="rounded,filled";
    color="#d4d4d8";
    fillcolor="#f8fafc";
    fontsize=20;
    t2 [
      fillcolor="#dcfce7",
      color="#16a34a",
      label=<
        <B>Task 2</B><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="18">Textbook alignment</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Curricular Alignment Architect</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Books -&gt; book skills + prerequisites</FONT>
      >
    ];
    t3 [
      fillcolor="#ffedd5",
      color="#ea580c",
      label=<
        <B>Task 3</B><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="18">Market alignment</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Market Demand Analyst</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Jobs -&gt; market skills</FONT>
      >
    ];
    { rank=same; t2; t3; }
  }

  subgraph cluster_application {
    label="Resource application";
    style="rounded,filled";
    color="#d4d4d8";
    fillcolor="#f8fafc";
    fontsize=20;
    t4 [
      fillcolor="#f3e8ff",
      color="#9333ea",
      label=<
        <B>Task 4</B><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="18">Resource curation</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Textual Resource Analyst + Video Agent</FONT><BR ALIGN="LEFT"/>
        <FONT POINT-SIZE="15">Skills -&gt; readings + videos</FONT>
      >
    ];
  }

  graphdb [
    fillcolor="#f8fafc",
    color="#64748b",
    label=<
      <B>Shared Neo4j curriculum knowledge graph</B><BR ALIGN="LEFT"/>
      <FONT POINT-SIZE="15">Unified coordination layer for chapters, concepts, book skills,</FONT><BR ALIGN="LEFT"/>
      <FONT POINT-SIZE="15">market skills, and downstream resources</FONT>
    >
  ];

  review [
    shape=note,
    style="filled",
    fillcolor="#fef3c7",
    color="#d97706",
    fontsize=16,
    label="Instructor review at textbook approval,\\nskill curation, and concept review"
  ];

  labtitle -> t1 [style=invis];
  t1 -> t2;
  t1 -> t3;
  t2 -> t4;
  t3 -> t4;

  t1 -> graphdb [style=dashed, color="#94a3b8", arrowhead=none];
  t2 -> graphdb [style=dashed, color="#94a3b8", arrowhead=none];
  t3 -> graphdb [style=dashed, color="#94a3b8", arrowhead=none];
  t4 -> graphdb [style=dashed, color="#94a3b8", arrowhead=none];

  review -> t2 [style=dotted, color="#d97706", arrowhead=none];
  review -> t3 [style=dotted, color="#d97706", arrowhead=none];

  { rank=same; review; }
}
"""

    path.parent.mkdir(parents=True, exist_ok=True)
    dot_path = path.with_suffix(".dot")
    dot_path.write_text(dot_source.strip() + "\n")
    subprocess.run(
        [dot, "-Gdpi=220", "-Tpng", str(dot_path), "-o", str(path)],
        check=True,
    )


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    style: str = "Body Text",
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    paragraph.text = text
    paragraph.style = style
    if alignment is not None:
        paragraph.alignment = alignment


def format_table(table: Table) -> None:
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.9), Inches(2.1), Inches(3.2), Inches(1.8)]
    headers = ["Task", "Main agent(s)", "Role in the framework", "Dependency"]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = widths[col]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        cell.paragraphs[0].style = "Body Text"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_idx, row in enumerate(TASK_TABLE_ROWS, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = value
            cell.paragraphs[0].style = "Body Text"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph starting with {startswith!r} not found")


def update_docx() -> Path:
    REVISION_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    OUTPUT_DOCX.chmod(0o644)
    generate_figure(FIGURE_PATH)

    doc = Document(OUTPUT_DOCX)

    image_rid = doc.inline_shapes[0]._inline.graphic.graphicData.pic.blipFill.blip.embed
    doc.part.related_parts[image_rid]._blob = FIGURE_PATH.read_bytes()

    overview = find_paragraph(doc, "Lab Tutor uses a course knowledge graph")
    task1 = find_paragraph(doc, "Task 1 (Document Extraction)")
    task2 = find_paragraph(doc, "Task 2 (Textbook Alignment)")
    task3 = find_paragraph(doc, "Task 3 (Market Alignment)")
    task4 = find_paragraph(doc, "Task 4 (Resource Curation)")
    task_dependency = find_paragraph(doc, "Tasks 2 and 3 are independent enrichment paths")
    figure1_caption = find_paragraph(doc, "Figure 1. Multi-agent system overview")

    set_paragraph_text(overview, FRAMEWORK_PARAGRAPHS[0])
    set_paragraph_text(task1, FRAMEWORK_PARAGRAPHS[1])
    set_paragraph_text(task2, FRAMEWORK_PARAGRAPHS[2])
    delete_paragraph(task3)
    delete_paragraph(task4)
    delete_paragraph(task_dependency)
    set_paragraph_text(figure1_caption, FIGURE_CAPTION, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    table_caption = insert_paragraph_after(figure1_caption, TABLE_1_CAPTION, style="Body Text")
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=len(TASK_TABLE_ROWS) + 1, cols=4)
    format_table(table)
    table_caption._p.addnext(table._tbl)

    task1_body = find_paragraph(doc, "Teacher-uploaded materials are parsed")
    task1_body.style = "Body Text"

    course_lead = find_paragraph(doc, "Table 1 summarizes results on 8 held-out Big Data jobs.")
    set_paragraph_text(course_lead, COURSE_TABLE_LEAD)
    course_caption = find_paragraph(doc, "Table 1 Course-level held-out alignment results.")
    set_paragraph_text(
        course_caption,
        COURSE_TABLE_CAPTION,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    course_discussion = find_paragraph(doc, "Market-skill enrichment produced the strongest result")
    set_paragraph_text(course_discussion, COURSE_DISCUSSION)

    student_lead = find_paragraph(doc, "Table 2 shows transfer from 7 student-selected jobs to 55 remaining jobs.")
    set_paragraph_text(student_lead, STUDENT_TABLE_LEAD)
    student_caption = find_paragraph(doc, "Table 2 Student-centered opportunity-spillover results.")
    set_paragraph_text(
        student_caption,
        STUDENT_TABLE_CAPTION,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    student_discussion = find_paragraph(doc, "Skills from 7 target jobs")
    set_paragraph_text(student_discussion, STUDENT_DISCUSSION)

    limitations = find_paragraph(doc, "Market-derived skills explain most measurable lift")
    set_paragraph_text(limitations, DISCUSSION_LIMITATIONS)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    result = update_docx()
    print(result)
