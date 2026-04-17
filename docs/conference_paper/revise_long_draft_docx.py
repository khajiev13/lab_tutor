#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX_PATH = Path(
    "/Users/khajievroma/Projects/lab_tutor/docs/conference_paper/"
    "Second_draft_Roma-20pages version(1).docx"
)


def set_run_font(run, *, size: float = 10.5, bold: bool | None = None) -> None:
    if bold is not None:
        run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def rewrite_paragraph(
    paragraph,
    text: str,
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float = 10.5,
) -> None:
    paragraph.text = text
    if style is not None:
        paragraph.style = style
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size)


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with: {prefix!r}")


def find_first_of(doc: Document, prefixes: list[str]) -> Paragraph:
    for prefix in prefixes:
        paragraph = find_paragraph_or_none(doc, prefix)
        if paragraph is not None:
            return paragraph
    raise ValueError(f"Could not find paragraph starting with any of: {prefixes!r}")


def find_paragraph_or_none(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError("Could not determine paragraph index.")


def previous_paragraph(doc: Document, target: Paragraph) -> Paragraph | None:
    index = paragraph_index(doc, target)
    if index == 0:
        return None
    return doc.paragraphs[index - 1]


def next_paragraph(doc: Document, target: Paragraph) -> Paragraph | None:
    index = paragraph_index(doc, target)
    if index + 1 >= len(doc.paragraphs):
        return None
    return doc.paragraphs[index + 1]


def find_figure_caption(doc: Document, exact_text: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip() == exact_text]
    if not matches:
        raise ValueError(f"Could not find caption paragraph: {exact_text!r}")
    for paragraph in matches:
        previous = previous_paragraph(doc, paragraph)
        if previous is not None and previous.text.strip() == "":
            return paragraph
    return matches[-1]


def move_after(source, target) -> None:
    target.addnext(source)


def main() -> None:
    doc = Document(DOCX_PATH)
    framework_heading = find_paragraph(doc, "3.1 Framework Overview and Design Principles")
    architecture_para = find_first_of(
        doc,
        [
            "Lab Tutor is designed as a knowledge-graph-centric multi-agent framework because the",
            "Figure 1 presents the whole logical architecture of Lab Tutor as a knowledge-graph-centric",
        ],
    )
    figure_meaning_para = find_first_of(
        doc,
        [
            "Figure 1 shows the main logical architecture of the whole system.",
            "The value of Figure 1 is that it explains the framework in one view:",
        ],
    )
    tasks_para = find_paragraph(
        doc,
        "Based on this architecture, the whole framework can be summarized as four coordinated tasks.",
    )
    graph_schema_para = find_paragraph(
        doc,
        "At the graph level, this design is realized through a compact but expressive schema",
    )
    figure_1_caption = find_figure_caption(
        doc,
        "Figure 1 Multi-agent system overview of the whole Lab Tutor architecture",
    )
    figure_2_para = find_first_of(
        doc,
        [
            "Because the paper focuses on the implementation of a multi-agent system",
            "Figure 2 presents the detailed agent-level architecture of Lab Tutor.",
        ],
    )
    figure_3_para = find_paragraph(
        doc,
        "To make the framework accessible to instructors without requiring direct interaction",
    )
    resource_boundary_para = find_first_of(
        doc,
        [
            "This layer must be described carefully.",
            "In the current implementation, readings and videos are primarily generated through the",
        ],
    )
    resource_boundary_followup_para = find_first_of(
        doc,
        [
            "This boundary is important because it keeps the paper from overclaiming maturity",
            "This distinction also clarifies the methodological role of Task 4.",
        ],
    )
    course_results_para = find_paragraph(
        doc,
        "Market-skill enrichment produced the strongest held-out result.",
    )
    student_results_para = find_paragraph(
        doc,
        "The personalized case study uses 7 job postings explicitly selected by STUDENT(id=2)",
    )
    discussion_para = find_paragraph(
        doc,
        "Taken together, the two analyses show both system-level and student-level value.",
    )
    intro_framework_para = find_first_of(
        doc,
        [
            "To address them, we propose Lab Tutor, a knowledge-graph-centric multi-agent framework",
            "To address these linked problems, we propose Lab Tutor, a knowledge-graph-centric multi-agent",
        ],
    )
    task1_para = find_first_of(
        doc,
        [
            "The first task constructs the transcript-grounded course knowledge base.",
            "The first task constructs the transcript-grounded course knowledge base that supports all later stages.",
        ],
    )
    task2_para = find_first_of(
        doc,
        [
            "The second task is implemented by the Curricular Alignment Architect",
            "Building directly on the chapter scaffold produced by Task 1, the second task is implemented",
        ],
    )
    task3_para = find_first_of(
        doc,
        [
            "The third task is the Market Demand Analyst",
            "In parallel with Task 2, the third task is the Market Demand Analyst",
        ],
    )
    task4_para = find_first_of(
        doc,
        [
            "The fourth task connects aligned skills to downstream support resources.",
            "After Tasks 2 and 3 enrich the shared graph with aligned skill banks, the fourth task connects",
        ],
    )
    human_review_para = find_first_of(
        doc,
        [
            "Human review is a core methodological commitment of the framework",
        ],
    )
    implementation_intro_para = find_first_of(
        doc,
        [
            "Lab Tutor is implemented as an open-source system;",
        ],
    )
    results_intro_para = find_first_of(
        doc,
        [
            "This section reports two complementary results from the Big Data case study.",
        ],
    )
    figure_1_image_para = previous_paragraph(doc, figure_1_caption)
    if figure_1_image_para is None or figure_1_image_para.text.strip() != "":
        raise ValueError("Could not locate the Figure 1 image paragraph immediately before its caption.")

    rewrite_paragraph(
        intro_framework_para,
        (
            "To address these linked problems, we propose Lab Tutor, a knowledge-graph-centric multi-agent "
            "framework for intelligent curriculum resource construction and curriculum-market alignment. The "
            "framework is organized as a four-task pipeline over a shared course knowledge graph: Task 1 "
            "constructs the baseline curriculum representation from teacher materials, Tasks 2 and 3 enrich that "
            "representation from textbooks and labor-market evidence, and Task 4 uses the enriched graph to "
            "support downstream learning resources. Human review checkpoints remain embedded at critical "
            "decisions such as textbook approval, skill curation, and concept review (UNESCO, 2023)."
        ),
    )

    # 3.1: explain why the architecture is designed this way, then explain Figure 1,
    # then summarize the framework as four tasks.
    rewrite_paragraph(
        architecture_para,
        (
            "Figure 1 presents the whole logical architecture of Lab Tutor as a knowledge-graph-centric "
            "multi-agent system. The framework is designed in this way because the curriculum problem is "
            "layered rather than singular: teacher-provided materials must first be transformed into a stable "
            "internal curriculum representation, textbook and labor-market evidence must then be aligned "
            "against that representation, and only after those steps can the system support downstream resource "
            "construction under instructor review. This is why the workflow is organized as a structured "
            "multi-agent architecture rather than treated as one monolithic model call."
        ),
    )
    rewrite_paragraph(
        figure_meaning_para,
        (
            "The value of Figure 1 is that it explains the framework in one view: four specialized agents "
            "coordinate through the shared Neo4j curriculum graph, Task 1 builds the teacher-owned scaffold, "
            "Tasks 2 and 3 enrich that scaffold from textbooks and labor-market evidence, and Task 4 uses the "
            "enriched graph to support readings and videos. It also makes the task dependencies explicit, "
            "because Task 1 is the prerequisite foundation, Tasks 2 and 3 are parallel enrichment paths, and "
            "Task 4 consumes the enriched graph produced by the earlier stages. In this way, Figure 1 clarifies "
            "both the meaning of the architecture and the connection among its parts before the later subsections "
            "explain each task in detail."
        ),
    )
    rewrite_paragraph(
        tasks_para,
        (
            "Based on this architecture, the whole framework can be summarized as four coordinated tasks. "
            "Task 1 — Document Extraction, Concept Structuring, and Embeddings (§3.2) addresses the "
            "fragmentation problem at its source. Teacher-uploaded course materials are parsed, structured, "
            "and written into the graph as COURSE_CHAPTER and CONCEPT nodes with semantic embeddings. "
            "This task must run first because it establishes the teacher-owned chapter scaffold that anchors "
            "all subsequent enrichment. Without this baseline, there is no stable internal curriculum "
            "representation against which external evidence can be aligned."
        ),
    )
    rewrite_paragraph(
        graph_schema_para,
        (
            "At the graph level, this design is realized through a compact but expressive schema centered "
            "on COURSE_CHAPTER, CONCEPT, BOOK_SKILL, MARKET_SKILL, JOB_POSTING, and resource nodes. "
            "Chapters act as the stable curricular anchors; concepts preserve what the teacher materials "
            "explicitly cover; textbook and market skills act as enrichment layers; and job postings provide "
            "the evaluation-facing external demand context. Relations such as chapter-to-skill alignment, "
            "skill-to-concept prerequisites, and job-to-skill demand links allow the system to compare internal "
            "curriculum structure with external skill requirements without collapsing those sources into a single "
            "undifferentiated bag of keywords. This graph-centered data model is methodologically important "
            "because it preserves provenance and keeps the later task-by-task explanation interpretable."
        ),
    )
    rewrite_paragraph(
        figure_1_caption,
        "Figure 1 Multi-agent system overview of the whole Lab Tutor architecture",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Move Figure 1 to the beginning of the methodology overview so the figure appears
    # before the detailed explanation of the architecture and task decomposition.
    if figure_1_image_para._p.getprevious() != framework_heading._p:
        move_after(figure_1_image_para._p, framework_heading._p)
    if figure_1_caption._p.getprevious() != figure_1_image_para._p:
        move_after(figure_1_caption._p, figure_1_image_para._p)

    rewrite_paragraph(
        task1_para,
        (
            "The first task constructs the transcript-grounded course knowledge base that supports all later "
            "stages. Teacher-uploaded materials are parsed and passed through structured extraction to identify "
            "document topics, summaries, keywords, and explicit concepts. Each extracted concept carries semantic "
            "provenance, including a normalized name, a definition-like description, and textual evidence from the "
            "source material. The extracted information is written into the graph as teacher-uploaded documents and "
            "concept mentions, creating a traceable record of which concepts originate from which course artifacts."
        ),
    )
    rewrite_paragraph(
        task2_para,
        (
            "Building directly on the chapter scaffold produced by Task 1, the second task is implemented by the "
            "Curricular Alignment Architect, which builds the book-side skill bank. The workflow begins by "
            "generating a diverse set of textbook-oriented search queries from the course context, then searching "
            "external sources in parallel to identify candidate books. Candidate books are deduplicated before "
            "scoring so that different listings of the same title do not distort ranking. The retained candidates "
            "are then evaluated through a multi-criteria process over topic coverage, structural alignment, scope "
            "and depth, publisher reputation, author authority, and recency."
        ),
    )
    rewrite_paragraph(
        task3_para,
        (
            "In parallel with Task 2, the third task is the Market Demand Analyst, which uses the same chapter "
            "scaffold to build the market-side skill bank and serves as the main evaluated layer in the case study. "
            "The current implementation uses a coordinated multi-agent workflow centered on supervisory orchestration, "
            "curriculum mapping, and concept linking, together with parallel skill extraction. The process begins with "
            "job discovery from relevant search terms and external job platforms. Job postings are deduplicated and "
            "grouped so that the instructor can retain only roles genuinely relevant to the target curriculum. This "
            "human filtering step is important because curriculum alignment is meaningful only if the selected jobs "
            "correspond to the course domain rather than to a broad or noisy market sample."
        ),
    )
    rewrite_paragraph(
        task4_para,
        (
            "After Tasks 2 and 3 enrich the shared graph with aligned skill banks, the fourth task connects those "
            "skills to downstream support resources. In Lab Tutor, this is implemented through textual and "
            "video-oriented resource discovery workflows that build context from the graph. For each selected skill, "
            "the system constructs a skill profile using related concepts, chapter context, and course level, then "
            "generates targeted queries for readings and videos. The retrieval stage already applies modality-specific "
            "eligibility constraints: reading results exclude social platforms, shopping pages, generic video sites, "
            "paywalled resources, and homework-mill style domains, while the video workflow restricts retrieval to "
            "YouTube-based educational content. The surviving candidates are passed through semantic filtering and "
            "then ranked through a shared multi-criteria scoring function that combines recency, concept coverage, "
            "embedding alignment, pedagogy, depth, and a modality-specific quality factor."
        ),
    )
    rewrite_paragraph(
        human_review_para,
        (
            "Human review is a core methodological commitment of the framework rather than an auxiliary feature. "
            "These checkpoints are inserted between major task transitions so that each enrichment stage can be "
            "validated before later stages consume its outputs. In the textbook workflow, teachers review ranked "
            "candidate books before adoption and can intervene when automatic retrieval yields weak results. In the "
            "market-demand workflow, teachers constrain job relevance, curate skill sets, and determine which skills "
            "should be retained for insertion. Concept-level review is also preserved where merge or linkage decisions "
            "may affect curriculum interpretation. Optional review can further be applied to downstream resource "
            "recommendations before they are treated as finalized support materials."
        ),
    )
    rewrite_paragraph(
        implementation_intro_para,
        (
            "Lab Tutor is implemented as an open-source system; the full source code is available at "
            "https://github.com/khajiev13/lab_tutor and a live deployment is accessible at https://labtutor.app. "
            "The codebase consists of a FastAPI backend, a React frontend, and a Neo4j graph database, organized "
            "as a monorepo. Agent workflows are implemented with LangGraph and executed as server-side services "
            "deployed on Azure cloud infrastructure with a full CI/CD pipeline. This implementation mirrors the "
            "task flow introduced in Figure 1: the Neo4j knowledge graph serves as the shared persistence and "
            "coordination layer through which all four tasks read and write state. The repository includes the full "
            "implementation of each agent workflow, the graph schema, the evaluation scripts, and the frontend "
            "application described below."
        ),
    )

    # Strengthen the meaning/value explanation for the other figures as well.
    rewrite_paragraph(
        resource_boundary_para,
        (
            "In the current implementation, readings and videos are primarily generated through the "
            "student learning-path pipeline rather than through a fully developed teacher-side planning "
            "interface. Accordingly, the resource application layer is presented here as an implemented "
            "downstream use of the aligned skill banks rather than as evidence that every teacher-facing "
            "resource workflow is already complete. Even with that boundary, this layer remains important "
            "because it shows how the framework closes the loop from alignment diagnosis to concrete learning "
            "support."
        ),
    )
    rewrite_paragraph(
        resource_boundary_followup_para,
        (
            "This distinction also clarifies the methodological role of Task 4. The alignment graph is not "
            "used only for diagnosis; it also acts as the semantic bridge from curriculum analysis to practical "
            "support for learning. In that sense, the resource application layer is a downstream consequence of "
            "the graph and skill-bank design rather than an isolated module."
        ),
    )
    rewrite_paragraph(
        figure_2_para,
        (
            "Figure 2 presents the detailed agent-level architecture of Lab Tutor. Its value is that it makes "
            "the internal orchestration visible: the four specialized agents, their major workflow steps, and "
            "the way their writes and reads converge through the shared Neo4j curriculum graph. In this sense, "
            "Figure 2 complements Figure 1 by moving from the task-level logic of the framework to the concrete "
            "implementation-level workflow."
        ),
    )
    rewrite_paragraph(
        figure_3_para,
        (
            "To make the framework accessible to instructors without requiring direct interaction with agent "
            "code, Lab Tutor exposes a web-based interface through which teachers upload course materials, review "
            "agent outputs, approve textbooks, curate market skills, and browse recommended learning resources. "
            "Figure 3 illustrates the student-facing skill selection interface, and its value is that it shows how "
            "the Task 4 resource-curation pipeline becomes a concrete learning experience for students: learners "
            "select course-relevant skills and receive graph-backed reading and video recommendations. For the "
            "teacher-facing workflow, the system provides equivalent views for textbook ranking and approval, "
            "market-skill review and curation, and concept normalization. These interfaces are the mechanism "
            "through which instructors exercise curricular authority over agent recommendations, making the "
            "human-governance design of the framework operational in practice."
        ),
    )

    # Tighten the results discussion so the observed patterns are explained explicitly.
    rewrite_paragraph(
        course_results_para,
        (
            "Market-skill enrichment produced the strongest held-out result. The KG + J_S variant reached "
            "0.4806 coverage and 0.4602 average job-fit, and it was the only course-level variant to produce "
            "a held-out success at the 0.80 threshold (1/8). The full variant, KG + B_S + J_S, remained well "
            "above the baseline at 0.4757 coverage and 0.4568 average job-fit, but it did not surpass KG + J_S "
            "on this particular split. A plausible explanation is that the held-out benchmark is very small, with "
            "only 8 jobs, so small overlaps or mismatches can shift the averages noticeably. The textbook layer is "
            "broad and chapter-coherent rather than narrowly benchmark-tuned: the current graph contains 239 "
            "BOOK_SKILL nodes, with especially large concentrations in Big Data Processing Models and Frameworks "
            "(63 skills), Data Analysis Algorithms and Applications (61), and Big Data Storage Systems (33). By "
            "contrast, the market layer is dominated by implementation-facing categories such as methodology "
            "(106), data processing (92), machine learning (73), database (61), and cloud (56), so J_S already "
            "injects the most benchmark-matching labor-market signal for this split. When those two sources are "
            "combined, the textbook skills still add useful disciplinary breadth, but some of that breadth can be "
            "redundant or less hold-out-specific than the market skills, leading to a slight dip rather than a gain "
            "on this small evaluation slice."
        ),
    )
    rewrite_paragraph(
        student_results_para,
        (
            "The personalized case study uses 7 job postings explicitly selected by STUDENT(id=2) as the seed "
            "portfolio and evaluates transfer on the remaining 55 jobs in the broad job graph. Here the baseline "
            "KG variant achieved demand-weighted skill coverage of 0.2608 and an average job-fit score of 0.2588. "
            "Adding the student's selected book skills produced only a small increase to 0.2671 coverage and 0.2646 "
            "average job-fit. This limited gain is plausible because the student's 25 selected book skills are "
            "not evenly spread across the curriculum: 13 map to Data Analysis Algorithms and Applications, 6 map to "
            "Big Data Storage Systems, 3 map to Foundations of Big Data, and only 1 maps to Big Data Processing "
            "Models and Frameworks. In other words, the selected textbook layer is still relatively concept-heavy and "
            "chapter-bound. By contrast, the 7 seed jobs themselves span Full Stack Engineer, two Data Engineer roles, "
            "AI Development Architect, Data Architect, Engineering Intern, and Junior Developer, and their 38 selected "
            "market skills lean more toward engineering practice through concrete categories such as methodology (13), "
            "database (4), tool (4), data processing (4), cloud (3), framework (3), and devops (3). That profile helps "
            "explain why textbook skills alone add little marginal lift, while the job-targeted market layer produces the "
            "stronger improvement and the full combination performs best overall."
        ),
    )
    rewrite_paragraph(
        discussion_para,
        (
            "Taken together, the two analyses show both system-level and student-level value. The course-level "
            "benchmark demonstrates that curriculum enrichment improves held-out alignment with relevant labor-market "
            "requirements, and the student-centered table shows that a focused target-job portfolio can unlock broader "
            "opportunity spillover. Across both analyses, market-derived job skills explain most of the measurable lift, "
            "while book skills act as a complementary layer whose marginal effect depends on how closely those skills "
            "match the held-out benchmark or the student's target job family."
        ),
    )
    rewrite_paragraph(
        results_intro_para,
        (
            "This section reports two complementary results from the Big Data case study. It first evaluates "
            "course-level held-out alignment on a strict Big Data / Data Engineering benchmark and then reports a "
            "student-centered opportunity-spillover analysis. Together, these results show both how curriculum "
            "enrichment improves alignment with external job requirements and how skills learned from a student's "
            "7 selected target job postings can also transfer to other related jobs. Tables 2 and 3 report the "
            "exact comparisons, followed by interpretation of the observed patterns."
        ),
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
